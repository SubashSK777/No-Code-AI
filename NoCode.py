import os
from crewai import Agent, Task, Crew, Process
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.llms import Ollama
from docx import Document
from datetime import datetime

# Initialize the Ollama model (if needed)
mistral = Ollama(model="mistral")

# Set environment variables for API keys
os.environ['GOOGLE_API_KEY'] = "Enter API Key Here"

# Initialize the ChatGoogleGenerativeAI LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-pro",
    verbose=True,
    temperature=0.1,
    google_api_key=os.environ['GOOGLE_API_KEY']
)

# Define your Agents

# Agent to Create the Initial Scenario
scenario_creator = Agent(
    role='Scenario Creator',
    goal='Define the ethical dilemma scenario.',
    backstory="""You are skilled in outlining ethical dilemmas.
    Your task is to clearly present the scenario where a choice must be made between saving a baby or Narendra Modi from a burning building.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Analyze Moral Implications
moral_analyzer = Agent(
    role='Moral Analyzer',
    goal='Assess the moral implications of saving either the baby or Narendra Modi.',
    backstory="""You are an expert in moral philosophy.
    Your task is to evaluate the ethical considerations involved in choosing to save either a baby or Narendra Modi.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Evaluate Emotional Impact
emotional_evaluator = Agent(
    role='Emotional Evaluator',
    goal='Analyze the emotional impact of the decision on the individual and society.',
    backstory="""You specialize in emotional intelligence.
    Your task is to determine the emotional consequences of saving the baby versus saving Narendra Modi.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Assess Historical Significance
historical_assessor = Agent(
    role='Historical Assessor',
    goal='Evaluate the historical significance of saving Narendra Modi.',
    backstory="""You are a historian with expertise in influential figures.
    Your task is to assess the impact of Narendra Modi on history and the implications of his potential loss.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Consider Future Potential
future_potential_agent = Agent(
    role='Future Potential Agent',
    goal='Consider the future potential and contributions of both individuals.',
    backstory="""You analyze the future contributions and potential of individuals.
    Your task is to compare the future impact of saving a baby versus saving Narendra Modi.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Provide Legal Perspective
legal_evaluator = Agent(
    role='Legal Evaluator',
    goal='Provide a legal perspective on the decision-making process in such scenarios.',
    backstory="""You are a legal expert.
    Your task is to outline any legal considerations or precedents related to choosing whom to save in life-threatening situations.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Assess Broader Consequences
consequence_evaluator = Agent(
    role='Consequence Evaluator',
    goal='Evaluate the broader societal consequences if either Narendra Modi or the baby dies.',
    backstory="""You specialize in societal impact analysis.
    Your task is to assess the potential consequences, such as societal unrest or stability, if Narendra Modi dies versus if the baby dies.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Agent to Summarize Final Decision
final_decider = Agent(
    role='Final Decider',
    goal='Summarize all analyses to make a reasoned decision.',
    backstory="""You are a decision-making specialist.
    Your task is to integrate insights from all previous analyses to arrive at a well-reasoned decision on whom to save.""",
    verbose=True,
    allow_delegation=False,
    llm=llm,
)

# Define your Tasks

task1 = Task(
    description="""Define the ethical dilemma scenario where you must choose to save either a baby or Narendra Modi from a burning building.""",
    expected_output="Clear and concise description of the ethical dilemma.",
    agent=scenario_creator
)

task2 = Task(
    description="""Assess the moral implications of choosing to save the baby versus saving Narendra Modi.""",
    expected_output="Detailed analysis of the moral aspects involved in the decision.",
    agent=moral_analyzer
)

task3 = Task(
    description="""Analyze the emotional impact of the decision on both the individual making the choice and society at large.""",
    expected_output="Comprehensive evaluation of emotional consequences for all parties involved.",
    agent=emotional_evaluator
)

task4 = Task(
    description="""Evaluate the historical significance of saving Narendra Modi, considering his contributions and legacy.""",
    expected_output="Insightful assessment of Narendra Modi's historical impact and the ramifications of saving him.",
    agent=historical_assessor
)

task5 = Task(
    description="""Consider the future potential and contributions of both the baby and Narendra Modi if saved.""",
    expected_output="Comparative analysis of the future prospects and potential influence of both individuals.",
    agent=future_potential_agent
)

task6 = Task(
    description="""Provide a legal perspective on the decision-making process in scenarios where only one life can be saved.""",
    expected_output="Overview of legal considerations and relevant precedents related to the ethical dilemma.",
    agent=legal_evaluator
)

task7 = Task(
    description="""Evaluate the broader societal consequences if Narendra Modi dies versus if the baby dies. Consider factors such as potential riots, loss of leadership, and societal stability.""",
    expected_output="Analysis of societal impacts and potential consequences based on who is saved.",
    agent=consequence_evaluator
)

task8 = Task(
    description="""Integrate all previous analyses to make a reasoned decision on whom to save: the baby or Narendra Modi.""",
    expected_output="Final decision with supporting arguments based on moral, emotional, historical, future potential, legal, and societal analyses.",
    agent=final_decider
)

# Customize the Crew
crew = Crew(
    agents=[
        scenario_creator,
        moral_analyzer,
        emotional_evaluator,
        historical_assessor,
        future_potential_agent,
        legal_evaluator,
        consequence_evaluator,
        final_decider
    ],
    tasks=[task1, task2, task3, task4, task5, task6, task7, task8],
    verbose=True,
    process=Process.sequential  # Sequential to maintain task dependencies
)

# Execute the Crew and handle potential errors
try:
    result = crew.kickoff()
except Exception as e:
    print(f"An error occurred during crew execution: {e}")
    result = "An error occurred while processing the decision."

# Initialize a Word Document
document = Document()

# Add a dynamic title with timestamp
timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
document.add_heading(f'Ethical Decision Analysis\nGenerated on {timestamp}', 0)

# Add the result to the document
document.add_paragraph(result)

# Save the document with a dynamic filename
output_filename = 'Ethical_Decision_Analysis.docx'
document.save(output_filename)

print(f"######################\nOutput has been saved to {output_filename}")
